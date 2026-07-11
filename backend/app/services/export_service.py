from uuid import UUID

from fastapi import HTTPException
from sqlalchemy.orm import Session

from app.models.gira import Gira
from app.models.inscricao_consulente import InscricaoConsulente

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter

from typing import List, Any
import csv
from io import BytesIO, StringIO

from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib.pagesizes import A4
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from dataclasses import dataclass


COLUMN_WIDTHS = {
    1: 10,
    2: 40,
    3: 22,
    4: 18,
    5: 18,
    6: 60,
}

def _buscar_gira(db: Session, gira_id: UUID, terreiro_id: UUID) -> Gira:
    gira = (
        db.query(Gira)
        .filter(
            Gira.id == gira_id,
            Gira.terreiro_id == terreiro_id,
            Gira.deleted_at.is_(None),
        )
        .first()
    )

    if not gira:
        raise HTTPException(status_code=404, detail="Gira não encontrada")

    return gira

def _buscar_inscricoes(db: Session, gira_id: UUID) -> list[InscricaoConsulente]:
    return (
        db.query(InscricaoConsulente)
        .filter(
            InscricaoConsulente.gira_id == gira_id,
            InscricaoConsulente.deleted_at.is_(None),
        )
        .order_by(InscricaoConsulente.posicao.asc())
        .all()
    )

def _montar_dados(gira: Gira, inscricoes: List[InscricaoConsulente]) -> dict[str, Any]:
    return {
        "gira": {
            "titulo": gira.titulo,
            "data": gira.data.strftime("%d/%m/%Y") if gira.data else "",
            "horario": str(gira.horario),
        },
        "inscricoes": [
            {
                "posicao": inscricao.posicao,
                "nome": (
                    inscricao.consulente.nome
                    if inscricao.consulente
                    else ""
                ),
                "telefone": (
                    inscricao.consulente.telefone
                    if inscricao.consulente
                    else ""
                ),
                "primeira_visita": (
                    "Sim"
                    if inscricao.consulente
                    and inscricao.consulente.primeira_visita
                    else "Não"
                ),
                "status": inscricao.status.value.replace("_", " ").title(),
                "observacoes": inscricao.observacoes or "",
            }
            for inscricao in inscricoes
        ],
    }

def export_excel(db: Session, gira_id: UUID, terreiro_id: UUID, ) -> BytesIO:

    gira = _buscar_gira(db, gira_id, terreiro_id)
    inscricoes = _buscar_inscricoes(db, gira.id)

    dados = _montar_dados(gira, inscricoes)
    
    wb = Workbook()
    assert wb.active is not None

    ws = wb.active
    ws.title = "Inscrições"

    # ----------------------------
    # Cabeçalho da Gira
    # ----------------------------

    ws["A1"] = "Gira"
    ws["B1"] = dados["gira"]["titulo"]

    ws["A2"] = "Data"

    ws["B2"] = dados["gira"]["data"]

    ws["A3"] = "Horário"
    ws["B3"] = dados["gira"]["horario"] 

    ws["A5"] = "Posição"
    ws["B5"] = "Nome"
    ws["C5"] = "Telefone"
    ws["D5"] = "Primeira visita"
    ws["E5"] = "Status"
    ws["F5"] = "Observações"

    gold_fill = PatternFill(
        fill_type="solid",
        fgColor="D4AF37",
    )

    white_font = Font(
        bold=True,
        color="FFFFFF",
    )

    thin = Side(style="thin")

    for cell in ws[5]:
        cell.fill = gold_fill
        cell.font = white_font
        cell.alignment = Alignment(horizontal="center")
        cell.border = Border(
            left=thin,
            right=thin,
            top=thin,
            bottom=thin,
        )

    row = 6

    for item in dados["inscricoes"]:

        ws.cell(row=row, column=1, value=item["posicao"])

        ws.cell(row=row, column=2, value=(item["nome"] if item["nome"] else "" ))

        ws.cell(row=row, column=3, value=(item["telefone"] if item["telefone"] else "" ))

        ws.cell(row=row, column=4, value=("Sim" if item["primeira_visita"] else "Não" ))

        status = item["status"] if item["status"] else ""

        ws.cell(row=row, column=5, value=status.replace("_", " ").title())

        ws.cell(row=row, column=6, value=item["observacoes"])

        row += 1

    for col, width in COLUMN_WIDTHS.items():
        ws.column_dimensions[get_column_letter(col)].width = width

    ws.freeze_panes = "A6"

    ws.auto_filter.ref = f"A5:F{row-1}"

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    return output

def export_csv(db: Session, gira_id: UUID, terreiro_id: UUID) -> BytesIO:
    gira = _buscar_gira(db, gira_id, terreiro_id)
    inscricoes = _buscar_inscricoes(db, gira.id)

    dados = _montar_dados(gira, inscricoes)

    output = StringIO()
    writer = csv.writer(output, delimiter=";")

    writer.writerow(["Gira", dados["gira"]["titulo"]])
    writer.writerow(["Data", dados["gira"]["data"]])
    writer.writerow(["Horário", dados["gira"]["horario"]])
    writer.writerow([])

    writer.writerow([
        "Posição",
        "Nome",
        "Telefone",
        "Primeira visita",
        "Status",
        "Observações",
    ])

    for inscricao in dados["inscricoes"]:
        writer.writerow([
            inscricao["posicao"],
            inscricao["nome"],
            inscricao["telefone"],
            inscricao["primeira_visita"],
            inscricao["status"],
            inscricao["observacoes"],
        ])

    buffer = BytesIO()
    buffer.write(output.getvalue().encode("utf-8-sig"))
    buffer.seek(0)

    return buffer

def export_pdf(db: Session, gira_id: UUID, terreiro_id: UUID) -> BytesIO:

    gira = _buscar_gira(db, gira_id, terreiro_id)
    inscricoes = _buscar_inscricoes(db, gira.id)

    dados = _montar_dados(gira, inscricoes)

    output = BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )

    styles = getSampleStyleSheet()

    titulo_style = styles["Title"]
    titulo_style.alignment = 1
    titulo_style.fontSize = 18
    titulo_style.spaceAfter = 12


    normal_style = styles["Normal"]
    normal_style.fontSize = 10
    
    cell_style = ParagraphStyle(
        "Cell",
        parent=styles["Normal"],
        fontSize=8,
        leading=10,
    )


    elementos = []


    # Título
    elementos.append(
        Paragraph(
            dados["gira"]["titulo"],
            titulo_style,
        )
    )


    # Informações da gira
    info = [
        [
            Paragraph("<b>Data</b>", normal_style),
            dados["gira"]["data"],
            Paragraph("<b>Horário</b>", normal_style),
            dados["gira"]["horario"],
        ]
    ]


    info_table = Table(
        info,
        colWidths=[
            2*cm,
            3*cm,
            2*cm,
            3*cm,
        ],
    )


    info_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#F5F5F5")),
                ("BOX", (0,0), (-1,-1), 0.3, colors.grey),
                ("INNERGRID", (0,0), (-1,-1), 0.3, colors.grey),
                ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ]
        )
    )


    elementos.append(info_table)

    elementos.append(
        Spacer(1, 0.8 * cm)
    )


    # Tabela inscrições

    tabela = [
        [
            Paragraph("<b>Pos.</b>", cell_style),
            Paragraph("<b>Nome</b>", cell_style),
            Paragraph("<b>Telefone</b>", cell_style),
            Paragraph("<b>1ª Visita</b>", cell_style),
            Paragraph("<b>Status</b>", cell_style),
            Paragraph("<b>Observações</b>", cell_style),
            Paragraph("<b>Consulta</b>", cell_style),
            Paragraph("<b>Passe</b>", cell_style),
            Paragraph("<b>Entidade</b>", cell_style),
        ]
    ]


    for item in dados["inscricoes"]:
        tabela.append(
            [
            Paragraph(str(item["posicao"]), cell_style),
            Paragraph(item["nome"], cell_style),
            Paragraph(item["telefone"], cell_style),
            Paragraph(item["primeira_visita"], cell_style),
            Paragraph(item["status"], cell_style),
            Paragraph(item["observacoes"] or "", cell_style),
            Paragraph("", cell_style),
            Paragraph("", cell_style),
            Paragraph("", cell_style),
        ]
        )


    table = Table(
    tabela,
    repeatRows=1,
    hAlign="LEFT",
)


    table.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0,0),
                    (-1,0),
                    colors.HexColor("#D4AF37")
                ),

                (
                    "TEXTCOLOR",
                    (0,0),
                    (-1,0),
                    colors.white
                ),

                (
                    "FONTNAME",
                    (0,0),
                    (-1,0),
                    "Helvetica-Bold"
                ),

                (
                    "FONTSIZE",
                    (0,0),
                    (-1,-1),
                    8
                ),

                (
                    "GRID",
                    (0,0),
                    (-1,-1),
                    0.25,
                    colors.grey
                ),

                (
                    "ALIGN",
                    (0,0),
                    (0,-1),
                    "CENTER"
                ),

                (
                    "VALIGN",
                    (0,0),
                    (-1,-1),
                    "MIDDLE"
                ),

                (
                    "BOTTOMPADDING",
                    (0,0),
                    (-1,0),
                    8
                ),

                (
                    "TOPPADDING",
                    (0,0),
                    (-1,0),
                    8
                ),
            ]
        )
    )


    elementos.append(table)


    doc.build(elementos)

    output.seek(0)

    return output

def export_docx(db: Session, gira_id: UUID, terreiro_id: UUID,) -> BytesIO:

    gira = _buscar_gira(db, gira_id, terreiro_id)
    inscricoes = _buscar_inscricoes(db, gira.id)

    dados = _montar_dados(gira, inscricoes)

    document = Document()

    titulo = document.add_heading(dados["gira"]["titulo"], level=1)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = document.add_paragraph()
    p.add_run("Data: ").bold = True
    p.add_run(dados["gira"]["data"])

    p = document.add_paragraph()
    p.add_run("Horário: ").bold = True
    p.add_run(dados["gira"]["horario"])

    document.add_paragraph()

    tabela = document.add_table(rows=1, cols=6)
    tabela.style = "Table Grid"

    cabecalho = tabela.rows[0].cells

    cabecalho[0].text = "Posição"
    cabecalho[1].text = "Nome"
    cabecalho[2].text = "Telefone"
    cabecalho[3].text = "1ª Visita"
    cabecalho[4].text = "Status"
    cabecalho[5].text = "Observações"

    for cell in cabecalho:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].font.bold = True
            paragraph.runs[0].font.size = Pt(10)

    for item in dados["inscricoes"]:
        cells = tabela.add_row().cells

        cells[0].text = str(item["posicao"])
        cells[1].text = item["nome"]
        cells[2].text = item["telefone"]
        cells[3].text = item["primeira_visita"]
        cells[4].text = item["status"]
        cells[5].text = item["observacoes"]

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output

@dataclass
class ExportFile:
    stream: BytesIO
    filename: str
    media_type: str
    
def export(db: Session, gira_id: UUID, terreiro_id: UUID, formato: str,):
    if formato == "xlsx":
        return ExportFile(
            stream=export_excel(db, gira_id, terreiro_id),
            filename="inscricoes.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    if formato == "csv":
        return ExportFile(
            stream=export_csv(db, gira_id, terreiro_id),
            filename="inscricoes.csv",
            media_type="text/csv",
        )

    if formato == "pdf":
        return ExportFile(
            stream=export_pdf(db, gira_id, terreiro_id),
            filename="inscricoes.pdf",
            media_type="application/pdf",
        )

    if formato == "docx":
        return ExportFile(
            stream=export_docx(db, gira_id, terreiro_id),
            filename="inscricoes.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    raise HTTPException(
        status_code=400,
        detail="Formato inválido.",
    )
    